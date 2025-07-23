
import pdfplumber
from docx import Document
from io import BytesIO

def extract_text(file):
    if file.name.endswith(".pdf"):
        with pdfplumber.open(file) as pdf:
            return "\n".join([p.extract_text() for p in pdf.pages if p.extract_text()])
    elif file.name.endswith(".docx"):
        doc = Document(file)
        return "\n".join([p.text for p in doc.paragraphs])
    else:
        return str(file.read(), "utf-8")

def generate_docx(text):
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output
